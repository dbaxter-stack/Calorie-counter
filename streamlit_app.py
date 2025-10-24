
from __future__ import annotations
from dataclasses import dataclass
from typing import Literal, Tuple
import streamlit as st
from io import BytesIO
from datetime import date
from docx import Document

st.set_page_config(page_title="9Round Pakenham â€¢ Calorie Calculator", page_icon="ðŸ¥Š", layout="centered")

st.markdown(
    """
    <style>
      .block-container { padding-top: 2rem; max-width: 980px; }
      .brand-title { font-weight: 900; font-size: 2.2rem; letter-spacing: 0.2px; color: #E4002B; }
      .sub-brand { color: #FFFFFF; font-weight: 700; }
      .kpi { border: 1px solid #2a2a2a; border-radius: 14px; padding: 16px 18px; background: #0B0B0B; }
      .kpi .value { font-size: 1.8rem; font-weight: 900; color: #FFFFFF; }
      .kpi .label { color: #D7DBE0; font-size: 0.95rem; text-transform: uppercase; letter-spacing: 0.6px; }
      .stButton>button { font-weight: 900; border-radius: 12px; text-transform: uppercase; background-color: #E4002B; color: white; }
      .footer { color: #D7DBE0; font-size: 0.9rem; margin-top: 1.25rem; text-align: center; }
      label, .stSelectbox label, .stNumberInput label, .stSlider label, .stTextInput label { text-transform: capitalize; }
    </style>
    """,
    unsafe_allow_html=True
)

ActivityKey = Literal["Sedentary", "Light", "Moderate", "Very", "Extra"]
BodyTypeKey = Literal["Ectomorph", "Mesomorph", "Endomorph"]
GoalKey = Literal["Maintain", "Lose Weight", "Gain Weight"]
SexKey = Literal["Male", "Female"]
UnitsKey = Literal["Metric", "Imperial"]

ACTIVITY_MULTIPLIER = {"Sedentary": 1.2, "Light": 1.375, "Moderate": 1.55, "Very": 1.725, "Extra": 1.9}
BODYTYPE_MULTIPLIER = {"Ectomorph": 1.05, "Mesomorph": 1.00, "Endomorph": 0.95}
KCAL_PER_KG = 7700.0

@dataclass
class CalcResult:
    bmr: float
    tdee: float
    calories: float
    daily_delta_kcal: float
    protein_g: float
    fats_g: float
    carbs_g: float

def mifflin_st_jeor(sex, age, weight_kg, height_cm):
    return 10 * weight_kg + 6.25 * height_cm - 5 * age + (5 if sex == "Male" else -161)

def convert_to_metric(units, weight, height):
    if units == "Imperial":
        return weight * 0.453592, height * 2.54
    return weight, height

def convert_mass_to_kg(units, mass):
    return mass * 0.453592 if units == "Imperial" else mass

def macro_split(total_cal, weight_kg, protein_g_per_kg=2.0, fat_percent=0.25):
    protein_g = protein_g_per_kg * weight_kg
    protein_kcal = protein_g * 4.0
    fat_kcal = total_cal * fat_percent
    fats_g = fat_kcal / 9.0
    carbs_kcal = max(total_cal - protein_kcal - fat_kcal, 0)
    carbs_g = carbs_kcal / 4.0
    return protein_g, fats_g, carbs_g

def calculate_intake(sex, age, weight, height, units, activity, body_type, goal, target_change_mass, target_weeks, protein_g_per_kg, fat_percent):
    w_kg, h_cm = convert_to_metric(units, weight, height)
    target_kg = convert_mass_to_kg(units, abs(target_change_mass))
    bmr = mifflin_st_jeor(sex, age, w_kg, h_cm)
    tdee = bmr * ACTIVITY_MULTIPLIER[activity] * BODYTYPE_MULTIPLIER[body_type]
    daily_delta = (target_kg * KCAL_PER_KG) / (target_weeks * 7.0) if target_weeks > 0 else 0.0
    if goal == "Lose Weight":
        daily_delta = -daily_delta
    elif goal == "Maintain":
        daily_delta = 0.0
    calories = tdee + daily_delta
    protein_g, fats_g, carbs_g = macro_split(calories, w_kg, protein_g_per_kg, fat_percent)
    return CalcResult(bmr, tdee, calories, daily_delta, protein_g, fats_g, carbs_g)

# Header Branding
st.markdown("<div class='brand-title'>9Round <span class='sub-brand'>Pakenham</span></div>", unsafe_allow_html=True)
st.caption("Estimate Your Daily Calorie Needs And Macronutrients Based On Activity, Body Type And Goal.")

# Body Type Descriptions
with st.expander("Body Type Descriptions", expanded=False):
    st.markdown(
        "- **Ectomorph** â€” Naturally slim; often finds it harder to gain weight/muscle; may benefit from a slightly higher intake.\n"
        "- **Mesomorph** â€” Athletic build; typically maintains more easily.\n"
        "- **Endomorph** â€” Stores body fat more readily; a slightly lower intake may help."
    )

with st.form("Calc Form"):
    c1, c2 = st.columns(2)
    with c1:
        client_name = st.text_input("Client Name", value="")
        sex = st.selectbox("Sex", ["Male", "Female"], index=0)
        age = st.number_input("Age (Years)", min_value=10, max_value=120, value=30, step=1)
        units = st.selectbox("Units", ["Metric", "Imperial"], index=0)
        weight = st.number_input("Weight (Kg Or Lb)", min_value=20.0, max_value=600.0, value=70.0, step=0.1)
        height = st.number_input("Height (Cm Or In)", min_value=80.0, max_value=300.0, value=175.0, step=0.1)
        activity = st.selectbox("Activity Level", ["Sedentary", "Light", "Moderate", "Very", "Extra"], index=2)
    with c2:
        body_type = st.selectbox("Body Type", ["Mesomorph", "Ectomorph", "Endomorph"], index=0)
        goal = st.selectbox("Goal", ["Maintain", "Lose Weight", "Gain Weight"], index=0,
                       help="Choose your objective. Maintain = stay around maintenance. Lose = create a moderate deficit. Gain = small surplus for muscle gain.")
        target_change, weeks = (0.0, 0)
        if goal != "Maintain":
            target_change = st.number_input("Target Weight Change (Kg Or Lb)", min_value=0.1, max_value=100.0, value=5.0, step=0.1,
                                help="How much weight you want to lose or gain in total.")
            weeks = st.number_input("Timeframe (Weeks)", min_value=1, max_value=104, value=8, step=1,
                         help="Shorter timeframes require larger daily calorie changes. Slower changes are usually more sustainable.")
        st.markdown("---")
        protein_g_per_kg = st.slider(
            "Protein (g/Kg)", 1.2, 2.6, 2.0, 0.1,
            help="Sets daily protein in grams per kilogram of body weight. Higher protein means fewer calories left for carbs and fats."
        )
        fat_percent = st.slider(
            "Fats (% Of Calories)", 0.20, 0.35, 0.25, 0.01,
            help="Controls the share of total calories from fats. Lower fat = more carbs; higher fat = fewer carbs."
        )
    submitted = st.form_submit_button("Calculate")

if submitted:
    res = calculate_intake(sex, age, weight, height, units, activity, body_type, goal, target_change, weeks, protein_g_per_kg, fat_percent)
    # ---- Safety Flags ----
    # Compute weekly rate of change in kg/week
    def to_kg(u, m):
        return m * 0.453592 if u == 'Imperial' else m
    weekly_rate_kg = 0.0
    if goal != 'Maintain' and weeks > 0:
        weekly_rate_kg = to_kg(units, abs(target_change)) / weeks
    # Thresholds (common guidance): loss > 0.75 kg/week, gain > 0.5 kg/week is aggressive
    if goal == 'Lose Weight' and weekly_rate_kg > 0.75:
        st.warning('This plan targets more than ~0.75 kg per week of weight loss. That is aggressive and may be hard to sustain. Consider extending your timeframe.')
    if goal == 'Gain Weight' and weekly_rate_kg > 0.5:
        st.warning('This plan targets more than ~0.5 kg per week of weight gain. That is aggressive and may be hard to sustain. Consider extending your timeframe.')
    # Also flag very large daily calorie changes relative to maintenance
    if res.tdee > 0 and abs(res.daily_delta_kcal) / res.tdee > 0.25:
        st.warning('Your required daily calorie change is greater than ~25% of maintenance. A smaller change is usually more realistic and healthier.')
    st.markdown(f"<h4 style='color:#fff;'>Target Daily Calories: {res.calories:.0f} kcal/day</h4>", unsafe_allow_html=True)
    st.write(f"**Protein:** {res.protein_g:.0f} g  â€¢  **Fats:** {res.fats_g:.0f} g  â€¢  **Carbs:** {res.carbs_g:.0f} g")

    def build_report():
        # Recompute safety flags for inclusion in the report
        def to_kg(u, m):
            return m * 0.453592 if u == 'Imperial' else m
        weekly_rate_kg = 0.0
        if goal != 'Maintain' and weeks > 0:
            weekly_rate_kg = to_kg(units, abs(target_change)) / weeks
        warn_loss = (goal == 'Lose Weight' and weekly_rate_kg > 0.75)
        warn_gain = (goal == 'Gain Weight' and weekly_rate_kg > 0.5)
        warn_delta = (res.tdee > 0 and abs(res.daily_delta_kcal) / res.tdee > 0.25)
        doc = Document()
        doc.add_heading("9Round Pakenham â€¢ Calorie & Macro Plan", level=1)
        doc.add_paragraph(f"Client: {client_name or '(Not Provided)'}    Date: {date.today().strftime('%d/%m/%Y')}")

        doc.add_heading("Goal", level=2)
        doc.add_paragraph(f"Objective: {goal}")
        if goal != "Maintain":
            doc.add_paragraph(f"Target Weight Change: {target_change} {'lb' if units=='Imperial' else 'kg'}")
            doc.add_paragraph(f"Timeframe: {weeks} weeks")
            doc.add_paragraph(f"Estimated Daily Calorie Adjustment: {res.daily_delta_kcal:+.0f} kcal/day")

        doc.add_heading("Energy & Activity Summary", level=2)
        doc.add_paragraph(f"Resting Energy Use: {res.bmr:.0f} kcal/day â€” Energy your body needs at rest to maintain vital functions like breathing and circulation.")
        doc.add_paragraph(f"Daily Maintenance Calories: {res.tdee:.0f} kcal/day â€” Total energy burned in a typical day including rest, movement and workouts.")
        doc.add_paragraph("Examples of activities that contribute to your TDEE:")
        doc.add_paragraph("â€¢ Walking 8,000â€“10,000 steps (â‰ˆ 300â€“400 kcal)")
        doc.add_paragraph("â€¢ 30 minutes of moderate cardio (â‰ˆ 250 kcal)")
        doc.add_paragraph("â€¢ 45-minute 9Round session or HIIT (â‰ˆ 400â€“500 kcal)")
        doc.add_paragraph("â€¢ Daily chores and errands (â‰ˆ 500â€“800 kcal)")
        doc.add_paragraph(f"Target Daily Calories: {res.calories:.0f} kcal/day â€” The estimated intake to help reach your goal.")

        # Warnings section (only if triggered)
        if warn_loss or warn_gain or warn_delta:
            doc.add_heading("Warnings", level=2)
            if warn_loss:
                doc.add_paragraph("This plan targets more than ~0.75 kg per week of weight loss. This is aggressive and may be hard to sustain.")
            if warn_gain:
                doc.add_paragraph("This plan targets more than ~0.5 kg per week of weight gain. This is aggressive and may be hard to sustain.")
            if warn_delta:
                doc.add_paragraph("The required daily calorie change exceeds ~25% of maintenance. A smaller change is usually more realistic and healthier.")
        doc.add_heading("Macronutrient Targets", level=2)
        doc.add_paragraph(f"Protein: {res.protein_g:.0f} g/day")
        doc.add_paragraph(f"Fats: {res.fats_g:.0f} g/day")
        doc.add_paragraph(f"Carbohydrates: {res.carbs_g:.0f} g/day")

        doc.add_heading("Notes", level=2)
        doc.add_paragraph("These figures are estimates only â€” individual energy needs can vary based on genetics, body composition, sleep, stress, and training intensity.")
        doc.add_paragraph("Use this plan as a starting point and adjust your intake based on changes in weight, energy, and performance. Consistency is key to progress.")

        p = doc.add_paragraph("")
        p.alignment = 1
        run = p.add_run("9Round Pakenham")
        run.bold = True

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio

    st.download_button("Download Report (.docx)", data=build_report(), file_name=f"9Round_Pakenham_Report_{date.today()}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Enter Your Details And Press **Calculate** To See Your Results.")
